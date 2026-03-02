import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import io
from datetime import datetime

from docx import Document
from docx.shared import Inches

st.set_page_config(page_title="NEST360 Internal DQA", layout="wide")

# =========================
# PASSWORD PROTECTION
# =========================
def check_password():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False

    if st.session_state.authenticated:
        return True

    st.title("🔐 NEST360 Internal DQA System")
    pw = st.text_input("Enter access password", type="password")

    if st.button("Login"):
        if pw == st.secrets.get("APP_PASSWORD", ""):
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Incorrect password.")
    st.stop()

check_password()

with st.sidebar:
    st.header("Controls")
    if st.button("Logout"):
        st.session_state.authenticated = False
        st.rerun()

st.title("🏥 NEST360 Health Data Reporting & DQA")

uploaded_file = st.file_uploader("Upload CSV or Excel", type=["csv", "xlsx"])
if uploaded_file is None:
    st.info("Upload a CSV/Excel export from REDCap to begin.")
    st.stop()

# =========================
# LOAD DATA
# =========================
if uploaded_file.name.lower().endswith(".csv"):
    df = pd.read_csv(uploaded_file)
else:
    df = pd.read_excel(uploaded_file)

original_cols = list(df.columns)

st.subheader("Data Preview")
st.dataframe(df.head(20), use_container_width=True)

# =========================
# COLUMN MAPPING
# =========================
st.subheader("1) Column mapping (select correctly)")

c1, c2, c3 = st.columns(3)
with c1:
    facility_col = st.selectbox(
        "Facility column",
        options=original_cols,
        index=original_cols.index("Facility Name") if "Facility Name" in original_cols else 0
    )
with c2:
    bw_col = st.selectbox(
        "Birth weight (grams) column",
        options=original_cols,
        index=original_cols.index("Birth weight (grams):") if "Birth weight (grams):" in original_cols else 0
    )
with c3:
    ga_col = st.selectbox(
        "Gestational age (weeks) column",
        options=original_cols,
        index=original_cols.index("Weeks:") if "Weeks:" in original_cols else 0
    )

c4, c5, c6 = st.columns(3)
with c4:
    cpap_col = st.selectbox(
        "CPAP administered column",
        options=["(None)"] + original_cols,
        index=(["(None)"] + original_cols).index("CPAP Administered:") if "CPAP Administered:" in original_cols else 0
    )
with c5:
    kmc_col = st.selectbox(
        "KMC administered column",
        options=["(None)"] + original_cols,
        index=(["(None)"] + original_cols).index("KMC Administered:") if "KMC Administered:" in original_cols else 0
    )
with c6:
    outcome_col = st.selectbox(
        "Outcome/Mortality column (preferred: Newborn status at discharge)",
        options=["(None)"] + original_cols,
        index=(["(None)"] + original_cols).index("Newborn status at discharge:") if "Newborn status at discharge:" in original_cols else 0
    )

# Optional additional important fields (for completeness dashboards)
opt1, opt2 = st.columns(2)
with opt1:
    discharge_wt_col = st.selectbox(
        "Discharge weight column (optional)",
        options=["(None)"] + original_cols,
        index=0
    )
with opt2:
    admit_date_col = st.selectbox(
        "Admission date column (optional)",
        options=["(None)"] + original_cols,
        index=0
    )
disch_date_col = st.selectbox(
    "Discharge date column (optional)",
    options=["(None)"] + original_cols,
    index=0
)

# =========================
# PREP WORK
# =========================
work = df.copy()
work[facility_col] = work[facility_col].astype(str).str.strip()

# =========================
# FILTER: FINAL selector + diagnostics (ignore Prospective/Retrospective)
# =========================
st.subheader("2) Dataset filter (FINAL only)")

final_col = "Are you entering a BASELINE or FINAL dataset record?"

def norm(x) -> str:
    return str(x).strip().lower()

filter_notes = []

if final_col in work.columns:
    with st.expander("Show REDCap export values (diagnostics)", expanded=False):
        st.write("Baseline/Final values (counts):")
        st.dataframe(work[final_col].dropna().astype(str).str.strip().value_counts().to_frame("count"))

    final_vals = sorted(work[final_col].dropna().astype(str).str.strip().unique())

    default_final = [v for v in final_vals if "final" in norm(v) or norm(v) == "2"]
    final_keep = st.multiselect(
        "Select the value(s) that mean FINAL",
        options=final_vals,
        default=default_final if default_final else final_vals[:1]
    )

    if not final_keep:
        st.error("Please select at least one FINAL value to continue.")
        st.stop()

    before = len(work)
    work = work[work[final_col].astype(str).str.strip().isin(final_keep)].copy()
    removed = before - len(work)

    st.success(f"Filtered to FINAL only: {len(work):,} records (removed {removed:,}).")
    filter_notes.append(f"Final values included: {final_keep}")

    if len(work) == 0:
        st.error("No records remain after FINAL filtering. Adjust the selected FINAL value(s).")
        st.stop()
else:
    st.warning("FINAL column not found in this file. Filter not applied.")
    filter_notes.append("FINAL filter not applied (column missing).")

# Always note PR field is ignored (per your request)
filter_notes.append("Prospective/Retrospective field ignored due to known entry errors.")

# =========================
# Derived numeric fields + categories
# =========================
bw = pd.to_numeric(work[bw_col], errors="coerce")
ga = pd.to_numeric(work[ga_col], errors="coerce")

def bw_category(x):
    if pd.isna(x): return "Missing"
    if x < 1000: return "<1000g"
    if 1000 <= x <= 1499: return "1000–1499g"
    if 1500 <= x <= 2499: return "1500–2499g"
    return "≥2500g"

def ga_category(x):
    if pd.isna(x): return "Missing"
    if x < 28: return "<28 weeks"
    if 28 <= x < 32: return "28–<32 weeks"
    if 32 <= x < 37: return "32–<37 weeks"
    return "≥37 weeks"

work["bw_cat"] = bw.map(bw_category)
work["ga_cat"] = ga.map(ga_category)

# =========================
# ANALYSIS MODE (All vs Single facility)
# =========================
st.subheader("3) Analysis mode")

all_facilities = sorted([f for f in work[facility_col].dropna().unique() if str(f).strip().lower() != "nan"])

mode = st.radio(
    "Choose analysis scope",
    ["All facilities (aggregate)", "Single facility (facility report)"],
    horizontal=True
)

if mode == "Single facility (facility report)":
    selected_facility = st.selectbox("Select facility", options=all_facilities)
    data = work[work[facility_col] == selected_facility].copy()
    st.info(f"Scope: **{selected_facility}** (n={len(data):,})")
else:
    selected_facility = None
    data = work.copy()
    st.info(f"Scope: **All facilities** (n={len(data):,})")

# =========================
# Missingness: BLANK ONLY
# =========================
# Do not treat 'Not recorded/Not readable' as missing. Only blanks/NaN are missing.
NOT_MISSING_VALUES = {
    "not recorded",
    "not readable",
    "not record",
    "not read",
    "not_recorded/not_readable",
    "not recorded/not readable",
}

def is_blank_only(x) -> bool:
    if pd.isna(x):
        return True
    s = str(x).strip()
    if s == "":
        return True
    if s.lower() in NOT_MISSING_VALUES:
        return False
    return False

def blank_count(series: pd.Series) -> int:
    return int(series.map(is_blank_only).sum())

# =========================
# DQA SUMMARY (scoped)
# =========================
st.subheader("4) DQA summary (blank-only missingness)")

total_rows = len(data)
duplicates = int(data.duplicated().sum())

key_cols = [facility_col, bw_col, ga_col]
if cpap_col != "(None)":
    key_cols.append(cpap_col)
if kmc_col != "(None)":
    key_cols.append(kmc_col)
if outcome_col != "(None)":
    key_cols.append(outcome_col)

missing_key = pd.Series({col: blank_count(data[col]) for col in key_cols}).sort_values(ascending=False)
missing_key_total = int(missing_key.sum())
missing_key_df = missing_key.to_frame("blank_missing_count")

bw_scoped = pd.to_numeric(data[bw_col], errors="coerce")
ga_scoped = pd.to_numeric(data[ga_col], errors="coerce")

bw_out_of_range = int(((bw_scoped < 300) | (bw_scoped > 5500)).sum())
ga_out_of_range = int(((ga_scoped < 20) | (ga_scoped > 44)).sum())

error_points = missing_key_total + duplicates + bw_out_of_range + ga_out_of_range
dqa_score = max(0.0, 100.0 - (error_points / (total_rows + 1)) * 100.0)

if dqa_score >= 95:
    status = "🟢 GREEN"
elif dqa_score >= 80:
    status = "🟡 YELLOW"
else:
    status = "🔴 RED"

a, b, c, d = st.columns(4)
a.metric("Records", f"{total_rows:,}")
b.metric("Duplicates", f"{duplicates:,}")
c.metric("DQA Score", f"{dqa_score:.2f}%")
d.metric("Status", status)

st.markdown("### Blank-only missingness (key fields)")
st.dataframe(missing_key_df, use_container_width=True)

st.markdown("### Validity checks (key clinical fields)")
st.write(f"- Birth weight out of range (<300 or >5500g): **{bw_out_of_range:,}**")
st.write(f"- Gestational age out of range (<20 or >44 weeks): **{ga_out_of_range:,}**")

# =========================
# Core breakdown tabs (BW / GA / Mortality / Completeness)
# =========================
st.subheader("5) Key breakdowns")

def yes_rate(series: pd.Series) -> float:
    s = series.astype(str).str.strip().str.lower()
    return float((s == "yes").mean() * 100)

def death_rate(series: pd.Series) -> float:
    s = series.astype(str).str.strip().str.lower()
    return float((s == "dead").mean() * 100)

tab_bw, tab_ga, tab_mort, tab_comp = st.tabs(
    ["Birth weight", "Gestational age", "Mortality", "Completeness & Validity"]
)

# ---- Birth weight tab
with tab_bw:
    st.markdown("### Birth weight categories (counts)")
    bw_counts = data["bw_cat"].value_counts()
    st.dataframe(bw_counts.to_frame("count"), use_container_width=True)

    bw_fig = plt.figure()
    bw_counts.plot(kind="bar")
    plt.xticks(rotation=45, ha="right")
    plt.ylabel("Count")
    st.pyplot(bw_fig)

    st.markdown("### Outcomes / interventions by birth weight category")
    grp = data.groupby("bw_cat", dropna=False)
    summary_bw = pd.DataFrame({"n": grp.size()})

    if cpap_col != "(None)":
        summary_bw["CPAP yes (%)"] = grp[cpap_col].apply(yes_rate)
    if kmc_col != "(None)":
        summary_bw["KMC yes (%)"] = grp[kmc_col].apply(yes_rate)
    if outcome_col != "(None)":
        summary_bw["Death (%)"] = grp[outcome_col].apply(death_rate)

    st.dataframe(summary_bw, use_container_width=True)

# ---- GA tab
with tab_ga:
    st.markdown("### Gestational age categories (counts)")
    ga_counts = data["ga_cat"].value_counts()
    st.dataframe(ga_counts.to_frame("count"), use_container_width=True)

    ga_fig = plt.figure()
    ga_counts.plot(kind="bar")
    plt.xticks(rotation=45, ha="right")
    plt.ylabel("Count")
    st.pyplot(ga_fig)

    st.markdown("### Outcomes / interventions by gestational age category")
    grp = data.groupby("ga_cat", dropna=False)
    summary_ga = pd.DataFrame({"n": grp.size()})

    if cpap_col != "(None)":
        summary_ga["CPAP yes (%)"] = grp[cpap_col].apply(yes_rate)
    if kmc_col != "(None)":
        summary_ga["KMC yes (%)"] = grp[kmc_col].apply(yes_rate)
    if outcome_col != "(None)":
        summary_ga["Death (%)"] = grp[outcome_col].apply(death_rate)

    st.dataframe(summary_ga, use_container_width=True)

# ---- Mortality tab
with tab_mort:
    if outcome_col == "(None)":
        st.warning("Select an Outcome/Mortality column to view mortality breakdowns.")
    else:
        st.markdown("### Mortality by facility (all facilities, FINAL filtered dataset)")
        tmp = work.copy()
        tmp_out = tmp[outcome_col].astype(str).str.strip().str.lower()
        tmp["dead_flag"] = (tmp_out == "dead").astype(int)

        mort_by_fac = tmp.groupby(facility_col).agg(
            records=("dead_flag", "size"),
            deaths=("dead_flag", "sum"),
        )
        mort_by_fac["Death (%)"] = (mort_by_fac["deaths"] / mort_by_fac["records"]) * 100
        mort_by_fac = mort_by_fac.sort_values("Death (%)", ascending=False)

        st.dataframe(mort_by_fac, use_container_width=True)

        st.markdown("### Top 10 highest mortality facilities")
        st.dataframe(mort_by_fac.head(10), use_container_width=True)

        st.markdown("### Bottom 10 lowest mortality facilities")
        st.dataframe(mort_by_fac.tail(10), use_container_width=True)

# ---- Completeness tab
with tab_comp:
    st.markdown("### Blank-only completeness (selected important fields)")
    important_cols = [bw_col, ga_col]
    if outcome_col != "(None)":
        important_cols.append(outcome_col)
    if cpap_col != "(None)":
        important_cols.append(cpap_col)
    if kmc_col != "(None)":
        important_cols.append(kmc_col)
    if discharge_wt_col != "(None)":
        important_cols.append(discharge_wt_col)
    if admit_date_col != "(None)":
        important_cols.append(admit_date_col)
    if disch_date_col != "(None)":
        important_cols.append(disch_date_col)

    comp = []
    for col in important_cols:
        blanks = blank_count(data[col])
        comp.append({
            "Field": col,
            "Records": total_rows,
            "Blank missing (n)": blanks,
            "Blank missing (%)": (blanks / (total_rows if total_rows else 1)) * 100
        })

    comp_df = pd.DataFrame(comp).sort_values("Blank missing (%)", ascending=False)
    st.dataframe(comp_df, use_container_width=True)

    st.markdown("### Additional validity checks (optional fields)")
    if discharge_wt_col != "(None)":
        dw = pd.to_numeric(data[discharge_wt_col], errors="coerce")
        dw_oob = int(((dw < 400) | (dw > 15000)).sum())
        st.write(f"- Discharge weight out of range (<400 or >15000g): **{dw_oob:,}**")

    if admit_date_col != "(None)" and disch_date_col != "(None)":
        ad = pd.to_datetime(data[admit_date_col], errors="coerce")
        dd = pd.to_datetime(data[disch_date_col], errors="coerce")
        bad_logic = int((dd < ad).sum())
        st.write(f"- Discharge date earlier than admission date: **{bad_logic:,}**")

# =========================
# Facility-level DQA table (always from full FINAL filtered dataset)
# =========================
st.subheader("6) Facility-level DQA (ranking)")

fac_table = work.groupby(facility_col).apply(
    lambda g: pd.Series({
        "records": len(g),
        "BW blank missing (%)": blank_count(g[bw_col]) / (len(g) if len(g) else 1) * 100,
        "GA blank missing (%)": blank_count(g[ga_col]) / (len(g) if len(g) else 1) * 100,
        "BW out-of-range (n)": int(((pd.to_numeric(g[bw_col], errors="coerce") < 300) |
                                   (pd.to_numeric(g[bw_col], errors="coerce") > 5500)).sum()),
        "GA out-of-range (n)": int(((pd.to_numeric(g[ga_col], errors="coerce") < 20) |
                                   (pd.to_numeric(g[ga_col], errors="coerce") > 44)).sum()),
    })
).sort_values("records", ascending=False)

st.dataframe(fac_table, use_container_width=True)

# =========================
# WORD REPORT HELPERS
# =========================
def fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=200)
    buf.seek(0)
    return buf

def df_to_docx_table(doc, df, title, max_rows=200):
    doc.add_heading(title, level=2)
    if df is None or df.empty:
        doc.add_paragraph("No data available.")
        return

    df2 = df.copy()
    if len(df2) > max_rows:
        df2 = df2.head(max_rows)
        doc.add_paragraph(f"(Showing first {max_rows} rows)")

    table = doc.add_table(rows=1, cols=len(df2.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df2.columns):
        hdr_cells[i].text = str(col)

    for _, row in df2.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = "" if pd.isna(val) else str(val)

def build_word_report(
    scope_label,
    dqa_score,
    status,
    total_rows,
    duplicates,
    missing_key_df,
    bw_counts_df,
    ga_counts_df,
    summary_bw,
    summary_ga,
    fac_table,
    bw_fig,
    ga_fig,
    filter_notes,
):
    doc = Document()
    doc.add_heading("NEST360 Neonatal DQA & Data Summary Report", level=0)

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Scope: {scope_label}")

    doc.add_heading("Filter notes", level=1)
    for n in filter_notes:
        doc.add_paragraph(f"- {n}")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        f"This report summarizes blank-only completeness, key validity checks, and core descriptive distributions. "
        f"DQA Score: {dqa_score:.2f}% ({status}). Records: {total_rows:,}. Duplicates: {duplicates:,}."
    )

    df_to_docx_table(
        doc,
        missing_key_df.reset_index().rename(columns={"index": "Field", "blank_missing_count": "Blank missing count"}),
        "Blank-only Missingness (Key Fields)"
    )

    df_to_docx_table(
        doc,
        bw_counts_df.reset_index().rename(columns={"index": "Birth weight category", "count": "Count"}),
        "Birth Weight Categories (Counts)"
    )

    df_to_docx_table(
        doc,
        ga_counts_df.reset_index().rename(columns={"index": "Gestational age category", "count": "Count"}),
        "Gestational Age Categories (Counts)"
    )

    doc.add_heading("Charts", level=1)
    doc.add_paragraph("Birth weight category distribution")
    doc.add_picture(fig_to_bytes(bw_fig), width=Inches(6.5))
    doc.add_paragraph("Gestational age category distribution")
    doc.add_picture(fig_to_bytes(ga_fig), width=Inches(6.5))

    df_to_docx_table(
        doc,
        summary_bw.reset_index().rename(columns={"bw_cat": "Birth weight category"}),
        "Interventions/Outcomes by Birth Weight Category"
    )
    df_to_docx_table(
        doc,
        summary_ga.reset_index().rename(columns={"ga_cat": "Gestational age category"}),
        "Interventions/Outcomes by Gestational Age Category"
    )
    df_to_docx_table(
        doc,
        fac_table.reset_index().rename(columns={fac_table.index.name or "index": "Facility"}),
        "Facility-level DQA Summary (All Facilities)",
        max_rows=500
    )

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# =========================
# DOWNLOAD OUTPUTS
# =========================
st.subheader("7) Download outputs")

# Excel filename
excel_name = "NEST360_DQA_Report.xlsx"
if mode == "Single facility (facility report)" and selected_facility:
    safe_name = "".join(ch for ch in selected_facility if ch.isalnum() or ch in [" ", "_", "-"]).strip()
    excel_name = f"NEST360_DQA_Report_{safe_name}.xlsx"

# Excel workbook (scoped)
excel_bytes = io.BytesIO()
with pd.ExcelWriter(excel_bytes, engine="xlsxwriter") as writer:
    data.to_excel(writer, index=False, sheet_name="data_with_categories")
    missing_key_df.to_excel(writer, sheet_name="blank_missing_key_fields")
    fac_table.to_excel(writer, sheet_name="facility_dqa")

    pd.DataFrame([{
        "scope": "single_facility" if mode == "Single facility (facility report)" else "all_facilities",
        "facility": selected_facility if mode == "Single facility (facility report)" else "ALL",
        "records": total_rows,
        "duplicates": duplicates,
        "blank_missing_key_total": missing_key_total,
        "bw_out_of_range": bw_out_of_range,
        "ga_out_of_range": ga_out_of_range,
        "dqa_score": round(dqa_score, 2),
        "status": status
    }]).to_excel(writer, index=False, sheet_name="summary")

st.download_button(
    "Download DQA workbook (Excel)",
    data=excel_bytes.getvalue(),
    file_name=excel_name,
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)

# Word report filename
docx_name = "NEST360_DQA_Report.docx"
if mode == "Single facility (facility report)" and selected_facility:
    safe_name = "".join(ch for ch in selected_facility if ch.isalnum() or ch in [" ", "_", "-"]).strip()
    docx_name = f"NEST360_DQA_Report_{safe_name}.docx"

scope_label = "ALL facilities (aggregate)" if mode == "All facilities (aggregate)" else f"Facility: {selected_facility}"

# Ensure charts exist even if user doesn't click into the tabs first
if "bw_fig" not in locals():
    tmp = data["bw_cat"].value_counts()
    bw_fig = plt.figure()
    tmp.plot(kind="bar")
    plt.xticks(rotation=45, ha="right")
    plt.ylabel("Count")

if "ga_fig" not in locals():
    tmp = data["ga_cat"].value_counts()
    ga_fig = plt.figure()
    tmp.plot(kind="bar")
    plt.xticks(rotation=45, ha="right")
    plt.ylabel("Count")

# Ensure summaries exist
if "summary_bw" not in locals():
    grp = data.groupby("bw_cat", dropna=False)
    summary_bw = pd.DataFrame({"n": grp.size()})
    if cpap_col != "(None)":
        summary_bw["CPAP yes (%)"] = grp[cpap_col].apply(yes_rate)
    if kmc_col != "(None)":
        summary_bw["KMC yes (%)"] = grp[kmc_col].apply(yes_rate)
    if outcome_col != "(None)":
        summary_bw["Death (%)"] = grp[outcome_col].apply(death_rate)

if "summary_ga" not in locals():
    grp = data.groupby("ga_cat", dropna=False)
    summary_ga = pd.DataFrame({"n": grp.size()})
    if cpap_col != "(None)":
        summary_ga["CPAP yes (%)"] = grp[cpap_col].apply(yes_rate)
    if kmc_col != "(None)":
        summary_ga["KMC yes (%)"] = grp[kmc_col].apply(yes_rate)
    if outcome_col != "(None)":
        summary_ga["Death (%)"] = grp[outcome_col].apply(death_rate)

word_bytes = build_word_report(
    scope_label=scope_label,
    dqa_score=dqa_score,
    status=status,
    total_rows=total_rows,
    duplicates=duplicates,
    missing_key_df=missing_key_df,
    bw_counts_df=data["bw_cat"].value_counts().to_frame("count"),
    ga_counts_df=data["ga_cat"].value_counts().to_frame("count"),
    summary_bw=summary_bw,
    summary_ga=summary_ga,
    fac_table=fac_table,
    bw_fig=bw_fig,
    ga_fig=ga_fig,
    filter_notes=filter_notes,
)

st.download_button(
    "Download report (Word .docx)",
    data=word_bytes.getvalue(),
    file_name=docx_name,
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
